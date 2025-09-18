import os
from flask import Flask, request, render_template_string, send_file, redirect, url_for, flash, jsonify, render_template
from werkzeug.utils import secure_filename
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from groq import Groq
from docx import Document
from PyPDF2 import PdfReader
from dotenv import load_dotenv
import re
import threading
import time
import contextlib
from nltk.tokenize import word_tokenize
from nltk.translate.meteor_score import meteor_score
import nltk
nltk.download('punkt', quiet=True)


# Load .env file
load_dotenv()

# Flask app
app = Flask(__name__)
app.secret_key = "supersecretkey"

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {"txt", "docx", "pdf"}

# Groq client
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY", "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"))

# Language codes (still kept for detection)
LANGUAGE_CODES = {
    'english': 'en', 'spanish': 'es', 'french': 'fr', 'german': 'de', 'italian': 'it',
    'portuguese': 'pt', 'russian': 'ru', 'chinese': 'zh', 'japanese': 'ja', 'korean': 'ko',
    'arabic': 'ar', 'hindi': 'hi', 'dutch': 'nl', 'swedish': 'sv', 'norwegian': 'no',
    'danish': 'da', 'finnish': 'fi', 'polish': 'pl', 'czech': 'cs', 'hungarian': 'hu',
    'turkish': 'tr', 'greek': 'el', 'hebrew': 'he', 'thai': 'th', 'vietnamese': 'vi',
    'indonesian': 'id', 'malay': 'ms', 'filipino': 'tl', 'ukrainian': 'uk', 'bulgarian': 'bg',
    'romanian': 'ro', 'croatian': 'hr', 'serbian': 'sr', 'slovenian': 'sl', 'slovak': 'sk',
    'estonian': 'et', 'latvian': 'lv', 'lithuanian': 'lt', 'catalan': 'ca', 'galician': 'gl'
}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def detect_language_with_groq(text):
    """Language detection using Groq LLM"""
    try:
        sample_text = text[:300].strip()
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": """Respond with ONLY the 2-letter ISO 639-1 code of the text language."""
                },
                {"role": "user", "content": sample_text}
            ],
            temperature=0.1,
            max_tokens=5
        )
        detected = response.choices[0].message.content.strip().lower()
        lang_code = ''.join(c for c in detected if c.isalpha())
        if len(lang_code) == 2:
            return lang_code
        return "en"
    except:
        return "en"

def translate_with_groq(text, src_lang, tgt_lang="en"):
    """Translate text using Groq LLM"""
    if not text.strip() or src_lang == tgt_lang:
        return text
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a professional translator. Translate the following text "
                        f"from {src_lang} to {tgt_lang} preserving meaning and formatting."
                    )
                },
                {"role": "user", "content": text}
            ],
            temperature=0.2,
            max_tokens=2048
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"[Translation failed: {e}]"

def translate_text_preserve_format(text, src_lang, tgt_lang="en"):
    """Translate line by line to preserve structure using Groq."""
    if src_lang == tgt_lang:
        return text
    lines = text.split('\n')
    translated_lines = []
    for line in lines:
        if line.strip():
            translated_lines.append(translate_with_groq(line, src_lang, tgt_lang))
        else:
            translated_lines.append('')
    return '\n'.join(translated_lines)

def read_file_content(filepath):
    ext = filepath.rsplit(".", 1)[1].lower()
    try:
        if ext == "txt":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif ext == "docx":
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == "pdf":
            reader = PdfReader(filepath)
            return "\n".join([p.extract_text() or '' for p in reader.pages])
        else:
            raise ValueError("Unsupported file type")
    except Exception as e:
        raise ValueError(f"Error reading file: {str(e)}")

from docx.shared import RGBColor

def save_translated_content(content, filename, src_lang, tgt_lang, original_ext):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = filename.rsplit(".", 1)[0]

    if original_ext == "txt":
        output_filename = f"{base_name}_{src_lang}_to_{tgt_lang}_{timestamp}.txt"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        return output_path, output_filename

    elif original_ext == "docx":
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        doc_in = Document(input_path)
        doc_out = Document()
        output_filename = f"{base_name}_{src_lang}_to_{tgt_lang}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        for para in doc_in.paragraphs:
            para_out = doc_out.add_paragraph()
            for run in para.runs:
                text = run.text.strip()
                if not text:
                    continue
                translated_text = translate_with_groq(text, src_lang, tgt_lang)
                new_run = para_out.add_run(translated_text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                if run.font.color and run.font.color.rgb:
                    new_run.font.color.rgb = RGBColor(
                        run.font.color.rgb[0],
                        run.font.color.rgb[1],
                        run.font.color.rgb[2]
                    )
        doc_out.save(output_path)
        return output_path, output_filename

    else:
        output_filename = f"{base_name}_{src_lang}_to_{tgt_lang}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(output_path)
        return output_path, output_filename

import sacrebleu
from rouge_score import rouge_scorer
from nltk.translate.meteor_score import meteor_score

def evaluate_translation(reference, hypothesis):
    metrics = {}
    bleu = sacrebleu.corpus_bleu([hypothesis], [[reference]])
    metrics['BLEU'] = round(bleu.score, 2)
    scorer = rouge_scorer.RougeScorer(['rougeL'], use_stemmer=True)
    rouge = scorer.score(reference, hypothesis)
    metrics['ROUGE-L'] = round(rouge['rougeL'].fmeasure * 100, 2)
    meteor = meteor_score([reference], hypothesis)
    metrics['METEOR'] = round(meteor * 100, 2)
    return metrics

@app.route("/", methods=["GET", "POST"])
def home():
    if request.method == "POST":
        input_type = request.form.get("input_type")
        tgt_lang = "en"
        if input_type == "file":
            if "file" not in request.files:
                flash("No file uploaded.", "error")
                return redirect(url_for("home"))

            file = request.files["file"]
            if file.filename == "":
                flash("No file selected.", "error")
                return redirect(url_for("home"))

            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                file.save(filepath)

                try:
                    start_time = datetime.now()
                    content = read_file_content(filepath)

                    if not content.strip():
                        flash("The uploaded file appears empty.", "error")
                        os.remove(filepath)
                        return redirect(url_for("home"))

                    src_lang = detect_language_with_groq(content)
                    if src_lang == tgt_lang:
                        translated_content = content
                        flash("File is already in English. No translation needed.", "success")
                    else:
                        translated_content = translate_text_preserve_format(content, src_lang, tgt_lang)
                        flash(f"File translated from {src_lang.upper()} to English successfully!", "success")

                    original_ext = filename.rsplit(".", 1)[1].lower()
                    output_path, output_filename = save_translated_content(
                        translated_content, filename, src_lang, tgt_lang, original_ext)

                                # ---- Translation ----
                    translated_text = translate_with_groq(original_ext, src_lang, tgt_lang)

                    # ---- ✅ Metrics with proper tokenization ----
                    hypothesis_tokens = word_tokenize(translated_text)
                    reference_tokens = word_tokenize(original_ext)


                    processing_time = (datetime.now() - start_time).total_seconds()
                    metrics = None
                    if src_lang != tgt_lang:
                        metrics = evaluate_translation(content, translated_content)

                    os.remove(filepath)
                    return render_template(
                        "indexf.html",
                        translated_text=translated_content[:2000] + "..." if len(translated_content) > 2000 else translated_content,
                        download_url=url_for("download_file", filename=output_filename),
                        detected_lang=src_lang,
                        processing_info=f"Translation completed in {processing_time:.2f} seconds. Language: {src_lang.upper()} → EN",
                        metrics=metrics
                    )
                except Exception as e:
                    if os.path.exists(filepath):
                        os.remove(filepath)
                    flash(f"Error processing file: {str(e)}", "error")
                    return redirect(url_for("home"))
            else:
                flash("Invalid file type. Upload .txt, .docx, or .pdf only.", "error")
                return redirect(url_for("home"))

    return render_template("indexf.html")

@app.route("/ajax_translate", methods=["POST"])
def ajax_translate():
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"translation": "", "detected_lang": ""})
        src_lang = detect_language_with_groq(text)
        if src_lang == "en":
            return jsonify({
                "translation": text,
                "detected_lang": src_lang,
                "message": "Text is already in English"
            })
        translated_text = translate_with_groq(text, src_lang, "en")
        return jsonify({
            "translation": translated_text,
            "detected_lang": src_lang
        })
    except Exception as e:
        return jsonify({
            "translation": "[Translation service temporarily unavailable]",
            "error": str(e)
        })

@app.route("/download/<filename>")
def download_file(filename):
    try:
        return send_file(
            os.path.join(OUTPUT_FOLDER, filename),
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        flash(f"Error downloading file: {str(e)}", "error")
        return redirect(url_for("home"))

if __name__ == "__main__":
    print("🚀 Starting Advanced AI Translator (Groq based)...")
    app.run(port="5080", debug=True)
